import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import os
import io
import plotly.express as px
from github import Github  
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import random

# --- Configuration ---
st.set_page_config(page_title="ICT Department Portal", layout="wide")
GITHUB_REPO = "motechbello1/ICT-Report-Portal" 

# File paths
ICT_DATA_FILE = "ict_master_log.csv"
EVENT_DATA_FILE = "ict_events_log.csv"
TRACKER_FILE = "last_report_date.txt"

# ==========================================
# DYNAMIC NARRATIVE ENGINE (DEEP PARAGRAPHS)
# ==========================================

def get_dynamic_intro(start, end, total, resolved, pending, rate, report_type="ICT"):
    if report_type == "ICT":
        intros = [
            f"During the comprehensive reporting window spanning {start} to {end}, the Information and Communication Technology (ICT) technical support division was engaged in a high volume of crucial diagnostic and remediation activities. The department actively managed {total} distinct support tickets encompassing network failures, hardware degradation, and software troubleshooting. The engineering team successfully drove a {rate:.1f}% operational resolution rate, fully closing {resolved} of these critical incidents. However, {pending} complex tasks remain in the active queue, requiring rollover prioritization in the immediate future to prevent cascading disruptions to daily organizational workflows.",
            f"This executive operational audit evaluates the structural health and responsiveness of the ICT department from {start} to {end}. Over this critical timeline, our infrastructure absorbed and processed {total} technical service requests across various organizational wings. Demonstrating strong technical agility, the support team achieved a {rate:.1f}% completion metric, successfully remediating {resolved} system-halting issues. The remaining {pending} unresolved items have been documented and escalated for advanced troubleshooting. Maintaining this high rate of resolution is paramount to ensuring that baseline administrative continuity is never compromised by technological bottlenecks.",
        ]
    elif report_type == "Event":
        intros = [
            f"Between the operational dates of {start} and {end}, the ICT logistics and deployment team engineered the complex technological frameworks required for {total} distinct organizational events. Crucially, the aggregate audience size that relied directly on our uninterrupted audio-visual, networking, and presentation infrastructure was estimated at {resolved} attendees. Successfully managing public-facing deployments of this scale requires extreme precision, rapid on-site troubleshooting, and exhaustive pre-event preparation to safeguard the organization's professional image.",
            f"This logistical impact report comprehensively audits the event-support timeframe from {start} to {end}. Our highly specialized field technicians were deployed to support the technological backbones of {total} scheduled gatherings. These deployments successfully facilitated seamless communication and presentation capabilities for a combined audience volume of approximately {resolved} individuals. The seamless execution of these events directly underscores the ICT department's capacity to handle high-stress, live-environment technological demands without critical failure.",
        ]
    elif report_type == "Master":
        intros = [
            f"This unified executive operational report merges the comprehensive data streams from both daily technical interventions and high-stakes logistical event coordination. Analyzing the period from {start} to {end}, the ICT department successfully managed a massive operational footprint consisting of {resolved} core diagnostic support tickets running concurrently alongside {pending} full-scale event infrastructure setups. By aggressively mitigating routine technological degradation while simultaneously facilitating major organizational gatherings, the ICT department has directly preserved massive amounts of institutional capital, productivity, and operational momentum.",
            f"Aggregating our cross-departmental databases yields a clear picture of the ICT team's macro-level operational health between {start} and {end}. The department operated at peak functional capacity, successfully executing {total} total discrete missions. This breaks down into {resolved} targeted departmental support interventions and the complex facilitation of {pending} major events. Operating successfully across these two highly demanding, parallel fronts proves the department's flexibility, though it places immense continuous pressure on our core engineering personnel and hardware reserves.",
        ]
    return random.choice(intros)

def get_dynamic_problem_analysis(top_prob, count):
    analyses = [
        f"A deep analytical dive into our hardware and software telemetry reveals that '{top_prob}' currently stands as our most pervasive and disruptive vulnerability, accounting for {count} separate emergency interventions during this cycle. A failure frequency of this magnitude strongly indicates an overarching systemic decay or network-wide vulnerability rather than isolated incidents of user error. Immediate managerial review is heavily recommended to determine if an aggressive hardware lifecycle replacement or network-wide patching protocol is required to permanently eliminate this repeating bottleneck.",
        f"Diagnostic trend tracking highlights '{top_prob}' as the primary technological disruption to our daily organizational workflows, triggering {count} distinct technical support tickets. When a single failure point demands this level of repetitive attention, it creates a massive drain on skilled labor hours that should be spent on proactive upgrades. We strongly advise addressing this root cause strategically—either through targeted capital expenditure on upgraded systems or rigorous vendor audits—to eradicate this continuous operational hemorrhage.",
    ]
    return random.choice(analyses)

def get_dynamic_staff_analysis(top_staff, type="support"):
    if type == "support":
        analyses = [
            f"Internally, human resource utilization metrics indicate that {top_staff} executed the absolute highest volume of specialized ticket resolutions. While this individual output is highly commendable and keeps the department afloat, it exposes a critical vulnerability: relying on a solitary 'lynchpin' technician creates a dangerous single point of failure. Management must ensure this operational output is sustainable and immediately implement aggressive cross-training protocols so that complex diagnostic skills are evenly distributed across the entire IT workforce.",
            f"Personnel tracking analytics show that {top_staff} absorbed the heaviest troubleshooting workload during this operational cycle, acting as the primary responder for the vast majority of critical incidents. Monitoring this human resource metric is essential not just for performance reviews, but for preventing severe staff burnout. Distributing this technical load across mid-level technicians remains a paramount operational priority for the upcoming quarter.",
        ]
    elif type == "event":
        analyses = [
            f"The human coordination metric clearly identifies {top_staff} at the absolute forefront of our event management and physical deployment strategies. Relying heavily on one primary coordinator for high-stress, public-facing, live-environment setups poses a severe organizational risk. Should this staff member be unavailable, institutional momentum will collapse. We must immediately mandate a shadowing program where junior technicians are forced to co-lead major deployments.",
            f"Logistical deployment data proves that {top_staff} anchored the highest number of complex physical event setups this week. Running continuous live events requires intense physical labor and high-pressure problem-solving. Expanding our active pool of heavily trained Audio-Visual coordinators is highly recommended to ensure we do not exhaust our premier field technicians."
        ]
    elif type == "master":
        analyses = [
            f"Visually mapping the unified workforce data establishes that {top_staff} absorbed the heaviest, most complex cross-disciplinary load, aggressively managing critical workflows across both micro-support tickets and macro-event coordination spaces. While heroic in the short term, relying on a solitary technician to bridge the gap between daily support and live events creates an incredibly dangerous operational bottleneck. Strategic Mandate: We must immediately initiate a forced cross-training protocol, systematically stripping {top_staff}'s routine duties and distributing them across junior personnel to flatten this dangerously skewed workload distribution.",
            f"The combined data matrix flags {top_staff} as the most heavily burdened staff member across all operational domains. This technician is effectively carrying the weight of multiple operational wings simultaneously. To prevent imminent staff burnout and the subsequent total collapse of our rapid-response capabilities, tasks must be aggressively redistributed. Management must prioritize hiring or promoting secondary specialists to directly shadow and offload responsibilities from our top performers."
        ]
    return random.choice(analyses)

def get_idle_dates_narrative(start_date, end_date, active_dates_series):
    active_dates = set(pd.to_datetime(active_dates_series).dt.date)
    delta = end_date - start_date
    all_dates = {start_date + timedelta(days=i) for i in range(delta.days + 1)}
    missing_dates = sorted(list(all_dates - active_dates))
    
    if not missing_dates:
        return "It is highly notable that operational tempo remained continuously aggressive; activities, emergencies, or deployments were actively recorded on every single calendar day within this reporting window. This indicates a heavily maximized operational environment."
    
    groups = []
    current_group = [missing_dates[0]]
    for i in range(1, len(missing_dates)):
        if (missing_dates[i] - missing_dates[i-1]).days == 1:
            current_group.append(missing_dates[i])
        else:
            groups.append(current_group)
            current_group = [missing_dates[i]]
    groups.append(current_group)
    
    text_parts = []
    for g in groups:
        if len(g) == 1:
            text_parts.append(g[0].strftime("%B %d, %Y"))
        else:
            text_parts.append(f"the period stretching from {g[0].strftime('%B %d')} to {g[-1].strftime('%B %d, %Y')}")
            
    date_str = ", ".join(text_parts)
    
    idle_phrases = [
        f"A chronological audit of the database reveals distinct operational lulls; specifically, absolutely zero emergency calls, technical interventions, or event setups were logged on {date_str}. Rather than viewing these as inactive days, management is strongly advised to officially designate these exact windows as 'Dark Days'—mandating preventative physical hardware cleaning, deep-level server audits, and automated firmware patching when organizational disruption risk is at its absolute lowest.",
        f"The telemetry reflects a complete pause in reactive support requirements on {date_str}. These documented idle periods represent vital, strategic opportunities that are currently being underutilized. Moving forward, the department should implement an automated protocol where any day experiencing zero incoming tickets immediately triggers a shift toward internal staff cross-training, inventory auditing, and cable management overhauls in high-traffic server environments."
    ]
    return random.choice(idle_phrases)

def get_strategic_recommendation():
    recs = [
        "Strategic Forward Outlook: Implementation of Strict Tiered SLAs (Service Level Agreements). Currently, all incoming support requests and event demands are treated with equal urgency, which fragments the focus of our engineering team. In the coming weeks, management must formalize a strict Tier 1 to Tier 3 emergency classification system. This will ensure that critical infrastructure failures are instantly escalated to senior engineers, while basic user errors are systematically routed to junior staff or automated self-help knowledge bases, vastly improving our overall resolution velocity.",
        "Strategic Forward Outlook: Aggressive Proactive Hardware Lifecycle Management. Our data continues to show repeated failure points on older physical systems. We must shift our operational philosophy from 'Reactive Troubleshooting' to 'Proactive Replacement.' It is the formal recommendation of the ICT department to execute a comprehensive audit of all organizational hardware deployed prior to 2021. Identifying and mass-replacing these aging machines before they experience catastrophic failure will drastically reduce our weekly ticket volume and permanently elevate organizational productivity.",
        "Strategic Forward Outlook: Cloud Migration and Redundancy Audits. As our operational and event data loads scale exponentially, relying solely on on-premise physical servers poses an unacceptable risk regarding data loss and hardware degradation. The department's primary goal for the upcoming operational phase must be evaluating a hybrid-cloud redundancy strategy. Shifting non-essential archiving and basic application hosting to secure cloud environments will immediately reduce the physical wear-and-tear on our local data centers and heavily insulate the organization against localized power or hardware failures.",
        "Strategic Forward Outlook: Standardization of High-Traffic Venue Infrastructure. Logistical data proves that setting up temporary audio-visual networks for recurring events consumes an immense amount of specialized labor hours. To reclaim this lost time, management must allocate capital to permanently wire and install standardized projector, PA, and networking systems in our most heavily utilized conference spaces. Transitioning these rooms to 'Plug-and-Play' environments will effectively eliminate hours of pre-event setup time, allowing ICT staff to focus on critical backend network stability."
    ]
    return random.choice(recs)

# --- GitHub Sync & Data Loaders ---
def push_to_github(file_path):
    try:
        if "GITHUB_TOKEN" in st.secrets:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo(GITHUB_REPO)
            with open(file_path, 'r') as file: content = file.read()
            try:
                contents = repo.get_contents(file_path)
                repo.update_file(contents.path, f"Auto-sync {file_path}", content, contents.sha)
                return True, "Synced to GitHub successfully!"
            except:
                repo.create_file(file_path, f"Auto-sync created {file_path}", content)
                return True, "Created and synced to GitHub successfully!"
        return False, "GITHUB_TOKEN not found. Saved locally."
    except Exception as e: return False, f"GitHub Sync Error: {str(e)}"

def load_and_clean_data(file_path, columns):
    if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
        df = pd.read_csv(file_path)
        for col in df.columns:
            if col != "Attendee Count": 
                df[col] = df[col].astype(str).str.strip().str.lower()
            else:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype(int)
        return df
    df = pd.DataFrame(columns=columns)
    df.to_csv(file_path, index=False)
    return df

def explode_staff_counts(df, column_name):
    if df.empty or column_name not in df.columns: return pd.Series(dtype=int)
    return df[column_name].dropna().astype(str).str.split(',').explode().str.strip().value_counts()

def get_last_report_date():
    if os.path.exists(TRACKER_FILE):
        with open(TRACKER_FILE, 'r') as f:
            try: return datetime.strptime(f.read().strip(), "%Y-%m-%d").date()
            except: pass
    return (datetime.now() - timedelta(days=7)).date()

def save_last_report_date(end_date):
    with open(TRACKER_FILE, 'w') as f: f.write(end_date.strftime("%Y-%m-%d"))

# ==========================================
# UI & NAVIGATION
# ==========================================
st.sidebar.title("🎛️ System Navigation")
app_mode = st.sidebar.radio("Select Portal Module:", ["🛠️ ICT Departmental Activities", "📅 ICT Events Dashboard", "📑 Master Combined Report"])
st.sidebar.markdown("---")

if "form_message" in st.session_state:
    if st.session_state.message_type == "success": st.success(st.session_state.form_message)
    else: st.warning(st.session_state.form_message)
    del st.session_state.form_message
    del st.session_state.message_type

# ==============================================================================
# MODULE 1: ICT DEPARTMENTAL ACTIVITIES
# ==============================================================================
if app_mode == "🛠️ ICT Departmental Activities":
    st.title("💻 Executive ICT Departmental Activities Portal")
    
    ict_columns = ["Date", "Day", "Time", "Department", "Reported By", "System ID", "Description", "Problem", "Action", "Parts", "IT Staff", "Status", "Remarks"]
    df_ict = load_and_clean_data(ICT_DATA_FILE, ict_columns)

    tab1, tab2, tab3, tab4 = st.tabs(["📝 Log Support Activity", "📊 Analytics Dashboard", "🗄️ Master Database", "📑 Executive Smart Report"])

    with tab1:
        with st.form("ict_log_form", clear_on_submit=True):
            t_col1, t_col2 = st.columns(2)
            with t_col1: log_date = st.date_input("Date", value=datetime.today())
            with t_col2: log_time = st.time_input("Time", value=datetime.now().time())
            
            col1, col2 = st.columns(2)
            with col1:
                department = st.text_input("Department")
                reported_by = st.text_input("Reported By")
                system_id = st.text_input("System ID")
                description = st.text_input("System Description")
                problem = st.text_area("Diagnosed Problem")
            with col2:
                action = st.text_area("Action Taken")
                parts = st.text_input("Replaced Parts (if any)")
                it_staff = st.text_input("IT Staff Name(s) - Separate with commas")
                status = st.selectbox("Status", ["resolved", "pending"]) 
                remarks = st.text_area("Remarks")
                
            if st.form_submit_button("💾 Save Activity Log & Sync"):
                new_data = {
                    "Date": log_date.strftime("%Y-%m-%d"), "Day": log_date.strftime("%A").lower(), "Time": log_time.strftime("%H:%M:%S"),
                    "Department": department.lower(), "Reported By": reported_by.lower(), "System ID": system_id.lower(), 
                    "Description": description.lower(), "Problem": problem.lower(), "Action": action.lower(),
                    "Parts": parts.lower(), "IT Staff": it_staff.lower(), "Status": status, "Remarks": remarks.lower()
                }
                temp_df = pd.concat([df_ict, pd.DataFrame([new_data])], ignore_index=True)
                temp_df = temp_df.sort_values(by=["Date", "Time"]).reset_index(drop=True)
                temp_df.to_csv(ICT_DATA_FILE, index=False)
                success, msg = push_to_github(ICT_DATA_FILE)
                st.session_state.message_type = "success" if success else "warning"
                st.session_state.form_message = msg
                st.rerun()

    with tab2:
        if not df_ict.empty:
            sort_by = st.selectbox("Analyze Metric:", ["Status", "Department", "IT Staff", "Problem"])
            col_name = "Problem" if sort_by == "Problem" else sort_by
            counts = explode_staff_counts(df_ict, col_name).reset_index() if col_name == "IT Staff" else df_ict[col_name].value_counts().reset_index()
            counts.columns = [col_name, "Count"]
            fig = px.pie(counts, values="Count", names=col_name) if col_name == "Status" else px.bar(counts, x=col_name, y="Count", color=col_name)
            st.plotly_chart(fig, use_container_width=True)

    with tab3:
        if not df_ict.empty: st.download_button("📥 Download CSV", df_ict.to_csv(index=False).encode('utf-8'), f"ICT_{datetime.now().strftime('%Y%m%d')}.csv", "text/csv")
        st.dataframe(df_ict, use_container_width=True)

    with tab4:
        st.subheader("Generate Dynamic Action Report")
        col_d1, col_d2 = st.columns(2)
        with col_d1: start_date = st.date_input("Start Date", value=get_last_report_date())
        with col_d2: end_date = st.date_input("End Date", value=get_last_report_date() + timedelta(days=7))

        if st.button("📄 Generate Report"):
            temp_df = df_ict.copy()
            temp_df['Date_Obj'] = pd.to_datetime(temp_df['Date'], errors='coerce').dt.date
            df_w = temp_df[(temp_df['Date_Obj'] >= start_date) & (temp_df['Date_Obj'] <= end_date)].copy()
            
            if df_w.empty: st.warning("No records found.")
            else:
                total_issues = len(df_w)
                resolved = len(df_w[df_w['Status'] == 'resolved'])
                pending = total_issues - resolved
                res_rate = (resolved / total_issues) * 100 if total_issues > 0 else 0
                
                doc = Document()
                doc.add_heading('Deep Executive ICT Operational Health Report', 0)
                doc.add_paragraph(f"Audited Timeline: {start_date} to {end_date}\nGenerated on: {datetime.now().strftime('%B %d, %Y')}")
                
                doc.add_heading('1. Executive Summary & Operational Tempo', level=1)
                doc.add_paragraph(get_dynamic_intro(start_date, end_date, total_issues, resolved, pending, res_rate, "ICT"))
                doc.add_paragraph(get_idle_dates_narrative(start_date, end_date, df_w['Date']))

                fig1, ax1 = plt.subplots(figsize=(4, 3))
                df_w['Status'].value_counts().plot(kind='pie', autopct='%1.1f%%', ax=ax1, colors=['#28a745', '#dc3545'])
                ax1.set_ylabel("")
                plt.title("Resolution Rate Analytics")
                fig1.savefig("t_stat.png", bbox_inches='tight')
                plt.close(fig1)
                
                fig2, ax2 = plt.subplots(figsize=(5, 3))
                prob_counts = df_w['Problem'].value_counts().head(5)
                prob_counts.plot(kind='barh', color='#ff6b6b', ax=ax2)
                ax2.invert_yaxis()
                plt.title("Systemic Technical Complaints")
                fig2.savefig("t_prob.png", bbox_inches='tight')
                plt.close(fig2)

                doc.add_heading('2. Diagnostic Breakdown & Vulnerabilities', level=1)
                doc.add_picture("t_stat.png", width=Inches(3.0))
                doc.add_picture("t_prob.png", width=Inches(4.5))
                doc.add_paragraph(get_dynamic_problem_analysis(prob_counts.index[0].title(), prob_counts.iloc[0]))

                fig3, ax3 = plt.subplots(figsize=(5, 3))
                df_w['Department'].value_counts().head(5).plot(kind='bar', color='#4dabf7', ax=ax3)
                plt.xticks(rotation=45, ha='right')
                plt.title("Departmental IT Demand Volume")
                fig3.savefig("t_dept.png", bbox_inches='tight')
                plt.close(fig3)

                doc.add_heading('3. Resource Allocation & Personnel Dynamics', level=1)
                doc.add_picture("t_dept.png", width=Inches(4.5))
                doc.add_paragraph(get_dynamic_staff_analysis(explode_staff_counts(df_w, "IT Staff").index[0].title(), "support"))

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                for f in ["t_stat.png", "t_prob.png", "t_dept.png"]: 
                    if os.path.exists(f): os.remove(f)
                save_last_report_date(end_date)
                
                st.success("✅ Deep Narrative Report Generated!")
                st.download_button("📥 Download Executive Document", doc_buffer, f"ICT_Deep_Report_{end_date}.docx")

# ==============================================================================
# MODULE 2: ICT EVENTS DASHBOARD
# ==============================================================================
elif app_mode == "📅 ICT Events Dashboard":
    st.title("🎟️ ICT Event Activity Dashboard")
    event_columns = ["Date", "Day", "Time", "Event Name", "Location", "Coordinator", "Equipment Deployed", "Attendee Count", "Status", "Remarks"]
    df_events = load_and_clean_data(EVENT_DATA_FILE, event_columns)

    etab1, etab2, etab3 = st.tabs(["📝 Log Event Activity", "🗄️ Database", "📑 Generate Event Report"])
    with etab1:
        with st.form("event_log_form", clear_on_submit=True):
            t_col1, t_col2 = st.columns(2)
            with t_col1: log_date = st.date_input("Date", value=datetime.today())
            with t_col2: log_time = st.time_input("Time", value=datetime.now().time())
            c1, c2 = st.columns(2)
            with c1:
                event_name = st.text_input("Event Name")
                location = st.text_input("Location")
                coordinator = st.text_input("Coordinator(s) - Comma Separated")
                equipment = st.text_area("Equipment Deployed")
            with c2:
                attendees = st.number_input("Attendees", min_value=0)
                status = st.selectbox("Status", ["planned", "completed", "cancelled"]) 
                remarks = st.text_area("Remarks")
            if st.form_submit_button("💾 Save Event"):
                new_event = {"Date": log_date.strftime("%Y-%m-%d"), "Day": log_date.strftime("%A").lower(), "Time": log_time.strftime("%H:%M:%S"), "Event Name": event_name.lower(), "Location": location.lower(), "Coordinator": coordinator.lower(), "Equipment Deployed": equipment.lower(), "Attendee Count": attendees, "Status": status, "Remarks": remarks.lower()}
                temp_df = pd.concat([df_events, pd.DataFrame([new_event])], ignore_index=True)
                temp_df.to_csv(EVENT_DATA_FILE, index=False)
                st.rerun()
                
    with etab2: st.dataframe(df_events, use_container_width=True)
    
    with etab3:
        st.subheader("Generate Narrative Event Report")
        col_d1, col_d2 = st.columns(2)
        with col_d1: start_date = st.date_input("Start Date", value=get_last_report_date())
        with col_d2: end_date = st.date_input("End Date", value=get_last_report_date() + timedelta(days=7))

        if st.button("📄 Generate Detailed Event Report"):
            temp_df = df_events.copy()
            temp_df['Date_Obj'] = pd.to_datetime(temp_df['Date'], errors='coerce').dt.date
            df_e = temp_df[(temp_df['Date_Obj'] >= start_date) & (temp_df['Date_Obj'] <= end_date)].copy()
            
            if not df_e.empty:
                doc = Document()
                doc.add_heading('Comprehensive ICT Event Logistics Report', 0)
                doc.add_paragraph(f"Audited Timeline: {start_date} to {end_date}\nGenerated on: {datetime.now().strftime('%B %d, %Y')}")
                
                doc.add_heading('1. Global Logistics & Impact Overview', level=1)
                doc.add_paragraph(get_dynamic_intro(start_date, end_date, len(df_e), df_e['Attendee Count'].sum(), 0, 0, "Event"))
                doc.add_paragraph(get_idle_dates_narrative(start_date, end_date, df_e['Date']))
                
                fig1, ax1 = plt.subplots(figsize=(5, 3))
                df_e['Location'].value_counts().plot(kind='bar', color='#17a2b8', ax=ax1)
                plt.xticks(rotation=45, ha='right')
                plt.title("Venue Utilization Map")
                fig1.savefig("e_loc.png", bbox_inches='tight')
                plt.close(fig1)

                fig2, ax2 = plt.subplots(figsize=(5, 3))
                explode_staff_counts(df_e, "Coordinator").plot(kind='pie', ax=ax2, autopct='%1.0f%%')
                ax2.set_ylabel("")
                plt.title("Coordinator Workload Distribution")
                fig2.savefig("e_coord.png", bbox_inches='tight')
                plt.close(fig2)

                doc.add_heading('2. Venue Strain & Personnel Execution Analysis', level=1)
                doc.add_picture("e_loc.png", width=Inches(4.5))
                doc.add_picture("e_coord.png", width=Inches(3.5))
                doc.add_paragraph(get_dynamic_staff_analysis(explode_staff_counts(df_e, "Coordinator").index[0].title(), "event"))

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                for f in ["e_loc.png", "e_coord.png"]: 
                    if os.path.exists(f): os.remove(f)
                
                st.success("✅ Deep Event Narrative Report Generated!")
                st.download_button("📥 Download Report", doc_buffer, f"Events_Deep_{end_date}.docx")

# ==============================================================================
# MODULE 3: MASTER COMBINED REPORT (MASSIVELY UPGRADED)
# ==============================================================================
elif app_mode == "📑 Master Combined Report":
    st.title("📑 Master ICT Consolidated Report")
    st.markdown("Generate a unified, highly descriptive executive document merging deep analytics, workload mapping, and actionable strategic recommendations.")
    
    ict_columns = ["Date", "Day", "Time", "Department", "Reported By", "System ID", "Description", "Problem", "Action", "Parts", "IT Staff", "Status", "Remarks"]
    event_columns = ["Date", "Day", "Time", "Event Name", "Location", "Coordinator", "Equipment Deployed", "Attendee Count", "Status", "Remarks"]
    df_ict = load_and_clean_data(ICT_DATA_FILE, ict_columns)
    df_events = load_and_clean_data(EVENT_DATA_FILE, event_columns)

    col_d1, col_d2 = st.columns(2)
    with col_d1: start_date = st.date_input("Start Date", value=get_last_report_date())
    with col_d2: end_date = st.date_input("End Date", value=get_last_report_date() + timedelta(days=7))

    if st.button("📄 Generate Deep Narrative Master Report"):
        df_ict['Date_Obj'] = pd.to_datetime(df_ict['Date'], errors='coerce').dt.date
        df_events['Date_Obj'] = pd.to_datetime(df_events['Date'], errors='coerce').dt.date
        
        rep_ict = df_ict[(df_ict['Date_Obj'] >= start_date) & (df_ict['Date_Obj'] <= end_date)]
        rep_evt = df_events[(df_events['Date_Obj'] >= start_date) & (df_events['Date_Obj'] <= end_date)]

        if rep_ict.empty and rep_evt.empty:
            st.error("No data found in either database for this timeline.")
        else:
            staff_act = explode_staff_counts(rep_ict, "IT Staff")
            staff_evt = explode_staff_counts(rep_evt, "Coordinator")
            combined_staff = pd.concat([staff_act, staff_evt], axis=1).fillna(0)
            combined_staff.columns = ['Activities', 'Events']
            combined_staff['Total'] = combined_staff['Activities'] + combined_staff['Events']
            combined_staff = combined_staff.sort_values(by='Total', ascending=False)
            
            total_ops = len(rep_ict) + len(rep_evt)
            
            doc = Document()
            doc.add_heading('Consolidated Executive ICT Operations & Strategic Insights', 0)
            doc.add_paragraph(f"Audited Timeline: {start_date} to {end_date}\nGenerated by Systems Intelligence on: {datetime.now().strftime('%B %d, %Y')}")
            
            doc.add_heading('Section 1: Macro-Operational Health & Timeline Analysis', level=1)
            doc.add_paragraph(get_dynamic_intro(start_date, end_date, total_ops, len(rep_ict), len(rep_evt), 0, "Master"))
            doc.add_paragraph(get_idle_dates_narrative(start_date, end_date, pd.concat([rep_ict['Date'], rep_evt['Date']])))

            # Tempo Graph
            tempo_df = pd.DataFrame({'Support': rep_ict['Date'].value_counts(),'Events': rep_evt['Date'].value_counts()}).fillna(0).sort_index()
            fig_line, ax_line = plt.subplots(figsize=(6, 3))
            tempo_df.plot(kind='line', marker='o', ax=ax_line, color=['#e83e8c', '#fd7e14'])
            plt.title("Daily Operational Tempo Matrix")
            plt.xticks(rotation=45, ha='right')
            fig_line.savefig("m_line.png", bbox_inches='tight')
            plt.close(fig_line)

            doc.add_heading('Section 2: Daily Workflow Pressure & Tempo', level=1)
            doc.add_picture("m_line.png", width=Inches(5.0))
            doc.add_paragraph("The trendline diagnostic mapped above tracks the exact volume of incoming support requests alongside scheduled logistical event dates. By mapping these daily operational pressure points, management can visually identify periods where the department is critically over-leveraged versus periods where resources can be safely redirected toward preventative maintenance protocols.")
            
            if not combined_staff.empty:
                fig_work, ax_work = plt.subplots(figsize=(6, 4))
                combined_staff[['Activities', 'Events']].head(7).plot(kind='bar', stacked=True, color=['#4dabf7', '#ffc107'], ax=ax_work)
                plt.title("Unified Cross-Departmental Personnel Workload")
                plt.xticks(rotation=45, ha='right')
                fig_work.savefig("m_work.png", bbox_inches='tight')
                plt.close(fig_work)

                doc.add_heading('Section 3: Human Resource Utilization & Bottleneck Tracking', level=1)
                doc.add_picture("m_work.png", width=Inches(5.0))
                doc.add_paragraph(get_dynamic_staff_analysis(combined_staff.index[0].title(), "master"))

            # NEW: STRATEGIC RECOMMENDATION ENGINE
            doc.add_heading('Section 4: Executive Strategic Directives', level=1)
            doc.add_paragraph(get_strategic_recommendation())

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            for f in ["m_work.png", "m_line.png"]: 
                if os.path.exists(f): os.remove(f)
            
            st.success("✅ Deep Narrative Master Report with Strategic Recommendations Generated!")
            st.download_button("📥 Download Master Report", doc_buffer, f"Master_Deep_{end_date}.docx")
