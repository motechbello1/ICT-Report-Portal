import pandas as pd
import os
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches

def generate_executive_report():
    data_file = "ict_master_log.csv"
    if not os.path.exists(data_file):
        print("Data file not found.")
        return

    df = pd.read_csv(data_file)
    df['Date'] = pd.to_datetime(df['Date'])
    
    # Filter for only the past 7 days
    one_week_ago = datetime.now() - timedelta(days=7)
    weekly_df = df[df['Date'] >= one_week_ago].copy()
    
    if weekly_df.empty:
        print("No records logged in the past 7 days.")
        return

    now = datetime.now()
    month_name = now.strftime("%B")
    week_num = f"Week_{(now.day - 1) // 7 + 1}"
    dir_path = os.path.join("Reports", month_name, week_num)
    os.makedirs(dir_path, exist_ok=True)

    # Calculate executive metrics
    total_issues = len(weekly_df)
    resolved = len(weekly_df[weekly_df['Status'] == 'Resolved'])
    pending = total_issues - resolved
    top_dept = weekly_df['Department'].mode()[0] if not weekly_df['Department'].empty else "N/A"
    top_staff = weekly_df['IT Staff'].mode()[0] if not weekly_df['IT Staff'].empty else "N/A"

    # --- Generate Chart Images ---
    chart_path = os.path.join(dir_path, "status_chart.png")
    plt.figure(figsize=(6, 4))
    weekly_df['Status'].value_counts().plot(kind='pie', autopct='%1.1f%%', colors=['#28a745', '#dc3545'])
    plt.title("Weekly Resolution Status")
    plt.ylabel("")
    plt.savefig(chart_path)
    plt.close()

    # --- Build the Word Document ---
    doc = Document()
    doc.add_heading('ICT Department - Executive Weekly Summary', 0)
    
    # Narrative Section
    doc.add_heading('1. Executive Narrative', level=1)
    narrative = (
        f"During the week ending {now.strftime('%Y-%m-%d')}, the ICT department managed a total of {total_issues} reported incidents. "
        f"The team successfully resolved {resolved} of these issues, leaving {pending} pending for follow-up. "
        f"The highest volume of requests originated from the {top_dept} department. "
        f"Special commendation to {top_staff} for leading the resolution metrics this week."
    )
    doc.add_paragraph(narrative)

    # Infographic Section
    doc.add_heading('2. Status Infographic', level=1)
    doc.add_picture(chart_path, width=Inches(5.0))
    os.remove(chart_path) # Clean up the image file after embedding

    # Table Section
    doc.add_heading('3. Complete Weekly Activity Log', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Define table headers (keeping it brief for Word doc width)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Dept'
    hdr_cells[2].text = 'System ID'
    hdr_cells[3].text = 'Problem'
    hdr_cells[4].text = 'Staff'
    hdr_cells[5].text = 'Status'

    # Fill table rows
    for index, row in weekly_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Date'].strftime('%Y-%m-%d'))
        row_cells[1].text = str(row['Department'])
        row_cells[2].text = str(row['System ID'])
        row_cells[3].text = str(row['Problem'])
        row_cells[4].text = str(row['IT Staff'])
        row_cells[5].text = str(row['Status'])

    doc_path = os.path.join(dir_path, f"Executive_Report_{now.strftime('%Y%m%d')}.docx")
    doc.save(doc_path)
    print(f"Executive Report and Table successfully generated at: {doc_path}")

if __name__ == "__main__":
    generate_executive_report()